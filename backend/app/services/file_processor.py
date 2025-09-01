"""
Phase 6: Enhanced File Processing Service

Comprehensive file processing with:
- Robust PDF/DOCX text extraction using multiple libraries
- Security validation and scanning
- File format validation with magic byte detection
- Preview generation
- Content analysis and metadata extraction
"""

import PyPDF2
import docx
import aiofiles
import hashlib
import mimetypes
import tempfile
from typing import Dict, Any, Optional, Tuple, List
import logging
import os
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass

# Enhanced PDF processing
try:
    import fitz  # PyMuPDF for better PDF handling
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logging.warning("PyMuPDF not available, falling back to PyPDF2 only")

# Image processing for previews
try:
    from PIL import Image, ImageDraw, ImageFont
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False
    logging.warning("Pillow not available, preview generation disabled")

# Security and validation
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False
    logging.warning("python-magic not available, using basic file validation")

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False
    logging.warning("chardet not available, using utf-8 encoding detection")

from fastapi import HTTPException

logger = logging.getLogger(__name__)

@dataclass
class ProcessingResult:
    """Enhanced processing result with security and metadata"""
    content: str
    metadata: Dict[str, Any]
    word_count: int
    extraction_method: str
    file_hash: str = ""
    preview_path: Optional[str] = None
    security_scan_passed: bool = True
    content_encoding: str = "utf-8"

class FileProcessor:
    """Enhanced file processor with security validation and content extraction"""
    
    def __init__(self, upload_dir: str = "uploads", max_file_size: int = 10 * 1024 * 1024):
        self.upload_dir = Path(upload_dir)
        self.max_file_size = max_file_size
        self.preview_dir = self.upload_dir / "previews"
        self.preview_dir.mkdir(exist_ok=True, parents=True)
        
        # Supported file types with their MIME types
        self.supported_types = {
            'pdf': ['application/pdf'],
            'docx': [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/vnd.ms-word.document.macroEnabled.12',
                'application/octet-stream'  # Sometimes browsers send this for .docx
            ],
            'txt': ['text/plain', 'text/txt'],
            'md': ['text/markdown', 'text/plain', 'text/x-markdown']
        }
        
        # Dangerous file signatures to block
        self.dangerous_signatures = {
            b'\x4D\x5A': 'Windows Executable',  # PE files
            b'\x7F\x45\x4C\x46': 'Linux Executable',  # ELF files
            b'\x89\x50\x4E\x47': 'PNG Image',  # Allow but validate
            b'\xFF\xD8\xFF': 'JPEG Image',  # Allow but validate
        }
    
    async def process_file(self, file_path: str, content_type: str, original_filename: str = "") -> ProcessingResult:
        """
        Enhanced file processing with security validation and content extraction
        
        Args:
            file_path: Path to the uploaded file
            content_type: MIME type of the file
            original_filename: Original filename for validation
            
        Returns:
            ProcessingResult with extracted content and metadata
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Read file content for validation
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            # Enhanced validation and security checks
            await self._validate_file_security(file_content, original_filename or file_path)
            
            # Route to appropriate processor based on content type
            if content_type == "application/pdf":
                result = await self._process_pdf_enhanced(file_path, file_content)
            elif "wordprocessingml" in content_type or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                result = await self._process_docx_enhanced(file_path, file_content)
            elif content_type.startswith("text/"):
                result = await self._process_text_enhanced(file_path, file_content)
            else:
                raise ValueError(f"Unsupported content type: {content_type}")
            
            # Generate file hash for deduplication
            result.file_hash = self._generate_file_hash(file_content)
            
            # Generate preview if possible
            if HAS_PILLOW:
                result.preview_path = await self._generate_preview(file_content, content_type, original_filename)
            
            return result
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return ProcessingResult(
                content="",
                metadata={"error": str(e)},
                word_count=0,
                extraction_method="error",
                security_scan_passed=False
            )
    
    async def _validate_file_security(self, content: bytes, filename: str) -> None:
        """Enhanced security validation for uploaded files"""
        # File size validation
        if len(content) > self.max_file_size:
            raise ValueError(f"File too large: {len(content) / (1024*1024):.1f}MB. Maximum: {self.max_file_size / (1024*1024):.1f}MB")
        
        if len(content) == 0:
            raise ValueError("Empty file not allowed")
        
        # File extension validation
        file_ext = Path(filename).suffix.lower().lstrip('.')
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: .{file_ext}")
        
        # MIME type validation using magic bytes if available
        if HAS_MAGIC:
            try:
                detected_mime = magic.from_buffer(content, mime=True)
                expected_mimes = self.supported_types[file_ext]
                
                # Special handling for text files with generic MIME types
                if file_ext in ['txt', 'md'] and detected_mime.startswith('text/'):
                    pass  # Allow any text MIME type
                elif detected_mime not in expected_mimes:
                    raise ValueError(f"MIME type mismatch: detected '{detected_mime}' for .{file_ext} file")
            except Exception as e:
                logger.warning(f"MIME detection failed: {e}")
        
        # Check for dangerous file signatures
        for dangerous_sig, description in self.dangerous_signatures.items():
            if content.startswith(dangerous_sig):
                if 'Executable' in description:
                    raise ValueError(f"Dangerous file signature detected: {description}")
        
        # Basic malicious pattern detection
        content_lower = content.lower()
        malicious_patterns = [
            b'<script',
            b'javascript:',
            b'vbscript:',
            b'onload=',
            b'onclick=',
            b'eval(',
            b'exec(',
        ]
        
        for pattern in malicious_patterns:
            if pattern in content_lower:
                logger.warning(f"Suspicious pattern detected in {filename}: {pattern}")
                # Note: In production, you might want to reject these files
                break
    
    async def _process_pdf_enhanced(self, file_path: str, content: bytes) -> ProcessingResult:
        """Enhanced PDF processing with multiple extraction methods"""
        extracted_text = ""
        metadata = {
            "pages": 0,
            "extraction_method": "hybrid",
            "has_images": False,
            "has_forms": False,
            "is_encrypted": False,
            "pdf_version": "",
            "creation_date": None,
            "modification_date": None,
            "creator": "",
            "producer": "",
            "title": "",
            "subject": "",
            "author": ""
        }
        
        try:
            # Primary extraction using PyMuPDF if available
            if HAS_PYMUPDF:
                try:
                    pdf_document = fitz.open(file_path)
                    
                    # Extract metadata
                    pdf_metadata = pdf_document.metadata
                    metadata.update({
                        "pages": pdf_document.page_count,
                        "pdf_version": f"{pdf_document.pdf_version()}",
                        "is_encrypted": pdf_document.needs_pass,
                        "title": pdf_metadata.get("title", ""),
                        "subject": pdf_metadata.get("subject", ""),
                        "author": pdf_metadata.get("author", ""),
                        "creator": pdf_metadata.get("creator", ""),
                        "producer": pdf_metadata.get("producer", ""),
                        "creation_date": pdf_metadata.get("creationDate", ""),
                        "modification_date": pdf_metadata.get("modDate", "")
                    })
                    
                    # Extract text from each page
                    text_parts = []
                    for page_num in range(pdf_document.page_count):
                        page = pdf_document[page_num]
                        page_text = page.get_text()
                        
                        if page_text.strip():
                            text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                        
                        # Check for images and forms
                        if page.get_images():
                            metadata["has_images"] = True
                        if page.get_widgets():
                            metadata["has_forms"] = True
                    
                    extracted_text = "\n\n".join(text_parts)
                    pdf_document.close()
                    
                except Exception as e:
                    logger.warning(f"PyMuPDF extraction failed: {e}")
                    extracted_text = ""
            
            # Fallback to PyPDF2 if PyMuPDF fails or is not available
            if not extracted_text.strip():
                logger.info("Using PyPDF2 for PDF extraction")
                extracted_text, pypdf2_metadata = await self._process_pdf_pypdf2(file_path)
                metadata.update(pypdf2_metadata)
                metadata["extraction_method"] = "pypdf2"
                
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Failed to extract PDF content: {str(e)}")
        
        if not extracted_text.strip():
            raise ValueError("PDF appears to be empty or contains only images/scanned content")
        
        word_count = len(extracted_text.split()) if extracted_text else 0
        
        return ProcessingResult(
            content=extracted_text.strip(),
            metadata=metadata,
            word_count=word_count,
            extraction_method=metadata["extraction_method"]
        )
    
    async def _process_pdf_pypdf2(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Fallback PDF extraction using PyPDF2"""
        content = ""
        metadata = {"extraction_method": "pypdf2"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata.update({
                    "pages": len(pdf_reader.pages),
                    "is_encrypted": pdf_reader.is_encrypted
                })
                
                # Extract PDF metadata
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata.update({
                        "title": pdf_meta.get("/Title", ""),
                        "author": pdf_meta.get("/Author", ""),
                        "subject": pdf_meta.get("/Subject", ""),
                        "creator": pdf_meta.get("/Creator", ""),
                        "producer": pdf_meta.get("/Producer", ""),
                        "creation_date": pdf_meta.get("/CreationDate", ""),
                        "modification_date": pdf_meta.get("/ModDate", "")
                    })
                
                # Extract text
                text_parts = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                content = "\n\n".join(text_parts)
                
        except Exception as e:
            logger.error(f"PyPDF2 extraction error: {e}")
            metadata["error"] = str(e)
        
        return content, metadata
    
    async def _process_docx_enhanced(self, file_path: str, content: bytes) -> ProcessingResult:
        """Enhanced DOCX processing with comprehensive metadata extraction"""
        extracted_text = ""
        metadata = {
            "paragraphs": 0,
            "tables": 0,
            "images": 0,
            "hyperlinks": 0,
            "word_count": 0,
            "character_count": 0,
            "has_headers_footers": False,
            "extraction_method": "python-docx",
            "docx_properties": {}
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract core properties
            props = doc.core_properties
            metadata["docx_properties"] = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
                "revision": str(props.revision) if props.revision else "",
                "category": props.category or "",
                "comments": props.comments or ""
            }
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            metadata["paragraphs"] = len(doc.paragraphs)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(" | ".join(row_data))
                
                if table_data:
                    table_text = "\n".join(table_data)
                    table_texts.append(f"[TABLE]\n{table_text}\n[/TABLE]")
            
            metadata["tables"] = len(doc.tables)
            
            # Combine all text
            all_text = "\n\n".join(text_parts)
            if table_texts:
                all_text += "\n\n" + "\n\n".join(table_texts)
            
            extracted_text = all_text
            
            # Count words and characters
            word_count = len(extracted_text.split()) if extracted_text else 0
            metadata["word_count"] = word_count
            metadata["character_count"] = len(extracted_text)
            
            # Check for headers/footers
            for section in doc.sections:
                if section.header.paragraphs or section.footer.paragraphs:
                    metadata["has_headers_footers"] = True
                    
                    # Extract header/footer text too
                    header_text = "\n".join([p.text for p in section.header.paragraphs if p.text.strip()])
                    footer_text = "\n".join([p.text for p in section.footer.paragraphs if p.text.strip()])
                    
                    if header_text:
                        extracted_text = f"[HEADER]\n{header_text}\n[/HEADER]\n\n{extracted_text}"
                    if footer_text:
                        extracted_text = f"{extracted_text}\n\n[FOOTER]\n{footer_text}\n[/FOOTER]"
                    break
                    
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Failed to extract DOCX content: {str(e)}")
        
        if not extracted_text.strip():
            raise ValueError("DOCX file appears to be empty")
        
        return ProcessingResult(
            content=extracted_text.strip(),
            metadata=metadata,
            word_count=word_count,
            extraction_method="python-docx"
        )
    
    async def _process_text_enhanced(self, file_path: str, content: bytes) -> ProcessingResult:
        """Enhanced text processing with encoding detection and analysis"""
        metadata = {
            "encoding": "utf-8",
            "confidence": 1.0,
            "line_count": 0,
            "word_count": 0,
            "character_count": 0,
            "extraction_method": "text",
            "file_size": len(content)
        }
        
        try:
            # Detect encoding if chardet is available
            if HAS_CHARDET:
                encoding_result = chardet.detect(content)
                detected_encoding = encoding_result.get('encoding', 'utf-8')
                confidence = encoding_result.get('confidence', 0.0)
                
                metadata.update({
                    "encoding": detected_encoding,
                    "confidence": confidence
                })
            else:
                detected_encoding = 'utf-8'
            
            # Decode content
            try:
                text_content = content.decode(detected_encoding)
                metadata["content_encoding"] = detected_encoding
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                text_content = content.decode('utf-8', errors='replace')
                metadata["encoding"] = "utf-8 (with replacements)"
                metadata["content_encoding"] = "utf-8"
            
            # Count statistics
            lines = text_content.split('\n')
            words = text_content.split()
            
            word_count = len(words)
            metadata.update({
                "line_count": len(lines),
                "word_count": word_count,
                "character_count": len(text_content)
            })
            
            # Basic content analysis
            if len(text_content) > 100:
                # Check if it looks like markdown
                markdown_indicators = ['#', '**', '*', '`', '[', ']', '(', ')']
                markdown_score = sum(1 for indicator in markdown_indicators if indicator in text_content[:500])
                
                if markdown_score > 3:
                    metadata["likely_format"] = "markdown"
                elif text_content.count('\t') > text_content.count(' ') * 0.1:
                    metadata["likely_format"] = "tab_separated"
                elif text_content.count(',') > len(lines) * 0.5:
                    metadata["likely_format"] = "csv_like"
                else:
                    metadata["likely_format"] = "plain_text"
            
            return ProcessingResult(
                content=text_content.strip(),
                metadata=metadata,
                word_count=word_count,
                extraction_method="text",
                content_encoding=metadata["content_encoding"]
            )
            
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            raise ValueError(f"Failed to extract text content: {str(e)}")
    
    async def _generate_preview(self, content: bytes, content_type: str, filename: str) -> Optional[str]:
        """Generate preview image for the document"""
        if not HAS_PILLOW:
            return None
            
        try:
            preview_filename = f"preview_{hashlib.md5(content).hexdigest()}.png"
            preview_path = self.preview_dir / preview_filename
            
            if content_type == 'application/pdf':
                return await self._generate_pdf_preview(content, preview_path)
            elif 'wordprocessingml' in content_type:
                return await self._generate_docx_preview(content, preview_path)
            elif content_type.startswith('text/'):
                return await self._generate_text_preview(content, preview_path, filename)
            
        except Exception as e:
            logger.warning(f"Preview generation failed for {filename}: {e}")
            # Preview generation failure shouldn't block file processing
            return None
        
        return None
    
    async def _generate_pdf_preview(self, content: bytes, preview_path: Path) -> Optional[str]:
        """Generate preview for PDF first page"""
        if not HAS_PYMUPDF:
            return None
            
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        try:
            pdf_document = fitz.open(temp_file_path)
            if pdf_document.page_count > 0:
                first_page = pdf_document[0]
                # Render page as image
                mat = fitz.Matrix(150/72, 150/72)  # 150 DPI
                pix = first_page.get_pixmap(matrix=mat)
                pix.save(str(preview_path))
                pdf_document.close()
                return str(preview_path.relative_to(self.upload_dir))
        except Exception as e:
            logger.warning(f"PDF preview generation failed: {e}")
        finally:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        
        return None
    
    async def _generate_docx_preview(self, content: bytes, preview_path: Path) -> Optional[str]:
        """Generate preview for DOCX (text-based preview)"""
        try:
            # For DOCX, create a text-based preview
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                doc = docx.Document(temp_file_path)
                text_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()][:10])  # First 10 paragraphs
                return await self._create_text_image_preview(text_content[:500], preview_path, "DOCX Document")
            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
        except Exception as e:
            logger.warning(f"DOCX preview generation failed: {e}")
            return None
    
    async def _generate_text_preview(self, content: bytes, preview_path: Path, filename: str) -> Optional[str]:
        """Generate preview for text files"""
        try:
            # Decode content
            if HAS_CHARDET:
                encoding_result = chardet.detect(content)
                encoding = encoding_result.get('encoding', 'utf-8')
            else:
                encoding = 'utf-8'
            
            try:
                text_content = content.decode(encoding)
            except UnicodeDecodeError:
                text_content = content.decode('utf-8', errors='replace')
            
            file_ext = Path(filename).suffix.upper()
            return await self._create_text_image_preview(text_content[:500], preview_path, f"{file_ext} Document")
        except Exception as e:
            logger.warning(f"Text preview generation failed: {e}")
            return None
    
    async def _create_text_image_preview(self, text: str, preview_path: Path, doc_type: str) -> Optional[str]:
        """Create an image preview from text content"""
        try:
            # Create a simple text preview image
            img = Image.new('RGB', (400, 300), color='white')
            draw = ImageDraw.Draw(img)
            
            try:
                # Try to use a nice font, fallback to default
                font = ImageFont.truetype("arial.ttf", 12)
                header_font = ImageFont.truetype("arial.ttf", 14)
            except:
                font = ImageFont.load_default()
                header_font = font
            
            # Draw header
            draw.text((10, 10), doc_type, fill='black', font=header_font)
            draw.line([(10, 30), (390, 30)], fill='gray', width=1)
            
            # Draw text content (word wrap)
            y_position = 45
            max_width = 380
            words = text.split()
            current_line = ""
            
            for word in words:
                test_line = f"{current_line} {word}".strip()
                try:
                    bbox = draw.textbbox((0, 0), test_line, font=font)
                    line_width = bbox[2] - bbox[0]
                except:
                    # Fallback for older Pillow versions
                    line_width = len(test_line) * 7  # Approximate
                
                if line_width <= max_width:
                    current_line = test_line
                else:
                    if current_line:
                        draw.text((10, y_position), current_line, fill='black', font=font)
                        y_position += 15
                        if y_position > 270:
                            break
                    current_line = word
            
            # Draw remaining text
            if current_line and y_position <= 270:
                draw.text((10, y_position), current_line, fill='black', font=font)
            
            img.save(preview_path, 'PNG')
            return str(preview_path.relative_to(self.upload_dir))
        except Exception as e:
            logger.warning(f"Text image preview creation failed: {e}")
            return None
    
    def _generate_file_hash(self, content: bytes) -> str:
        """Generate SHA-256 hash for file deduplication"""
        return hashlib.sha256(content).hexdigest()
    
    def validate_file_type(self, content_type: str, filename: str) -> bool:
        """Enhanced file type validation"""
        from app.core.config import settings
        
        # Get file extension
        file_ext = Path(filename).suffix.lower().lstrip('.')
        
        # Check if extension is allowed
        if file_ext not in settings.ALLOWED_EXTENSIONS:
            return False
        
        # Validate content type matches extension using our enhanced mapping
        valid_combinations = self.supported_types
        return content_type in valid_combinations.get(file_ext, [])

# Global instance
file_processor = FileProcessor()